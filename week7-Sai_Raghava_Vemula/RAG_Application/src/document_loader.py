"""
document_loader.py — Load and extract text from PDF, DOCX, and TXT files.

Each loader returns a list of ``langchain_core.documents.Document`` objects
with metadata including filename, page number, and source path.
"""

import os
from typing import List

from langchain_core.documents import Document

from src.logger import get_logger
from src.utils import detect_file_type, sanitize_text

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Individual loaders
# ---------------------------------------------------------------------------

def load_pdf(file_path: str) -> List[Document]:
    """
    Extract text from a PDF file using ``pypdf``.

    Skips pages that contain no meaningful text. Each page becomes
    one ``Document`` with metadata ``page_number`` and ``filename``.
    """
    from pypdf import PdfReader

    logger.info("Loading PDF: %s", file_path)
    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        logger.error("Failed to read PDF '%s': %s", file_path, exc)
        raise RuntimeError(f"Could not read PDF '{os.path.basename(file_path)}': {exc}") from exc

    documents: List[Document] = []
    filename = os.path.basename(file_path)

    for page_num, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        text = sanitize_text(raw_text)
        if not text:
            logger.debug("Skipping empty page %d in '%s'", page_num, filename)
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "filename": filename,
                    "page_number": page_num,
                    "source": file_path,
                },
            )
        )

    if not documents:
        logger.warning("PDF '%s' produced zero text pages.", filename)

    logger.info("Loaded %d page(s) from '%s'.", len(documents), filename)
    return documents


def load_docx(file_path: str) -> List[Document]:
    """
    Extract text from a DOCX file using ``python-docx``.

    All non-empty paragraphs are concatenated into a single ``Document``.
    """
    from docx import Document as DocxDocument

    logger.info("Loading DOCX: %s", file_path)
    try:
        doc = DocxDocument(file_path)
    except Exception as exc:
        logger.error("Failed to read DOCX '%s': %s", file_path, exc)
        raise RuntimeError(f"Could not read DOCX '{os.path.basename(file_path)}': {exc}") from exc

    paragraphs = [sanitize_text(p.text) for p in doc.paragraphs if sanitize_text(p.text)]
    filename = os.path.basename(file_path)

    if not paragraphs:
        logger.warning("DOCX '%s' contains no extractable text.", filename)
        return []

    full_text = "\n".join(paragraphs)
    documents = [
        Document(
            page_content=full_text,
            metadata={
                "filename": filename,
                "page_number": 1,
                "source": file_path,
            },
        )
    ]

    logger.info("Loaded DOCX '%s' (%d characters).", filename, len(full_text))
    return documents


def load_txt(file_path: str) -> List[Document]:
    """
    Read a plain-text file with UTF-8 encoding (BOM-safe).

    Returns a single ``Document``.
    """
    logger.info("Loading TXT: %s", file_path)
    filename = os.path.basename(file_path)

    try:
        with open(file_path, "r", encoding="utf-8-sig") as fh:
            raw_text = fh.read()
    except Exception as exc:
        logger.error("Failed to read TXT '%s': %s", file_path, exc)
        raise RuntimeError(f"Could not read TXT '{filename}': {exc}") from exc

    text = sanitize_text(raw_text)
    if not text:
        logger.warning("TXT '%s' is empty.", filename)
        return []

    documents = [
        Document(
            page_content=text,
            metadata={
                "filename": filename,
                "page_number": 1,
                "source": file_path,
            },
        )
    ]

    logger.info("Loaded TXT '%s' (%d characters).", filename, len(text))
    return documents


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
}


def load_document(file_path: str) -> List[Document]:
    """
    Auto-detect file type and extract text.

    Parameters
    ----------
    file_path : str
        Absolute path to the file.

    Returns
    -------
    List[Document]
        Extracted document pages / sections.

    Raises
    ------
    ValueError
        If the file type is unsupported.
    RuntimeError
        If the file cannot be read.
    """
    ext = detect_file_type(file_path)
    loader = _LOADERS[ext]
    return loader(file_path)


def load_multiple_documents(file_paths: List[str]) -> List[Document]:
    """
    Load and concatenate documents from multiple file paths.

    Files that fail to load are logged and skipped — they do not
    interrupt processing of the remaining files.
    """
    all_docs: List[Document] = []
    for path in file_paths:
        try:
            docs = load_document(path)
            all_docs.extend(docs)
        except Exception as exc:
            logger.error("Skipping file '%s': %s", path, exc)
    logger.info("Total documents loaded: %d from %d file(s).", len(all_docs), len(file_paths))
    return all_docs
